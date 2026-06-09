#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""文档格式转换服务。"""
import io
import logging
import re
import tempfile
from pathlib import Path
from typing import Tuple

logger = logging.getLogger(__name__)

_MAX_PAGES = 50
_MAX_FILE_BYTES = 20 * 1024 * 1024

SUPPORTED_CONVERSIONS: dict[str, set[str]] = {
    "pdf": {"docx", "md", "txt"},
    "docx": {"md", "txt"},
    "md": {"docx"},
    "txt": set(),
}

MIME_BY_EXT = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "md": "text/markdown",
    "txt": "text/plain",
}


class ConversionService:
    def convert(
            self,
            file_bytes: bytes,
            source_ext: str,
            target_ext: str,
            *,
            filename: str = "",
    ) -> Tuple[bytes, str, str]:
        if len(file_bytes) > _MAX_FILE_BYTES:
            raise ValueError("文件大小不能超过 20MB")

        source = source_ext.lower().lstrip(".")
        target = target_ext.lower().lstrip(".")
        allowed = SUPPORTED_CONVERSIONS.get(source, set())
        if target not in allowed:
            raise ValueError(f"暂不支持 {source} → {target} 转换")

        base_name = Path(filename or f"document.{source}").stem
        out_filename = f"{base_name}.{target}"

        if source == "pdf" and target == "docx":
            out_bytes = self._pdf_to_docx(file_bytes)
        elif source == "pdf" and target in {"md", "txt"}:
            out_bytes = self._pdf_to_text(file_bytes, as_markdown=target == "md").encode("utf-8")
        elif source == "docx" and target in {"md", "txt"}:
            out_bytes = self._docx_to_text(file_bytes, as_markdown=target == "md").encode("utf-8")
        elif source == "md" and target == "docx":
            text = self._decode_text(file_bytes)
            out_bytes = self._md_to_docx(text)
        else:
            raise ValueError(f"暂不支持 {source} → {target} 转换")

        mime = MIME_BY_EXT.get(target, "application/octet-stream")
        return out_bytes, mime, out_filename

    @staticmethod
    def _decode_text(data: bytes) -> str:
        for encoding in ("utf-8", "utf-8-sig", "gb18030", "latin-1"):
            try:
                return data.decode(encoding)
            except UnicodeDecodeError:
                continue
        return data.decode("utf-8", errors="replace")

    def _pdf_to_docx(self, data: bytes) -> bytes:
        try:
            from pdf2docx import Converter
        except ImportError as exc:
            raise ValueError("pdf2docx 未安装，无法执行 PDF 转 Word") from exc

        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as src:
            src.write(data)
            src_path = src.name
        out_path = src_path.replace(".pdf", ".docx")
        try:
            cv = Converter(src_path)
            try:
                cv.convert(out_path, start=0, end=min(_MAX_PAGES, 999))
            finally:
                cv.close()
            return Path(out_path).read_bytes()
        finally:
            Path(src_path).unlink(missing_ok=True)
            Path(out_path).unlink(missing_ok=True)

    def _pdf_to_text(self, data: bytes, *, as_markdown: bool) -> str:
        try:
            import fitz
        except ImportError as exc:
            raise ValueError("pymupdf 未安装，无法解析 PDF") from exc

        doc = fitz.open(stream=data, filetype="pdf")
        parts: list[str] = []
        try:
            for idx, page in enumerate(doc):
                if idx >= _MAX_PAGES:
                    parts.append("...(超出页数上限，已截断)")
                    break
                text = (page.get_text() or "").strip()
                if not text:
                    continue
                if as_markdown:
                    parts.append(f"## 第 {idx + 1} 页\n\n{text}")
                else:
                    parts.append(f"--- Page {idx + 1} ---\n{text}")
        finally:
            doc.close()
        return "\n\n".join(parts) if parts else ""

    def _docx_to_text(self, data: bytes, *, as_markdown: bool) -> str:
        try:
            from docx import Document
        except ImportError as exc:
            raise ValueError("python-docx 未安装，无法解析 Word 文档") from exc

        doc = Document(io.BytesIO(data))
        lines: list[str] = []
        for para in doc.paragraphs:
            text = (para.text or "").strip()
            if not text:
                continue
            style = (para.style.name or "").lower() if para.style else ""
            if as_markdown and "heading" in style:
                level = 1
                match = re.search(r"heading\s*(\d+)", style)
                if match:
                    level = min(int(match.group(1)), 6)
                lines.append(f"{'#' * level} {text}")
            elif as_markdown and "list" in style:
                lines.append(f"- {text}")
            else:
                lines.append(text)
        return "\n\n".join(lines)

    def _md_to_docx(self, markdown_text: str) -> bytes:
        try:
            import markdown as md_lib
            from docx import Document
            from docx.shared import Pt
        except ImportError as exc:
            raise ValueError("markdown/python-docx 未安装，无法生成 Word 文档") from exc

        html = md_lib.markdown(markdown_text or "")
        doc = Document()
        plain_lines = self._html_to_lines(html)
        if not plain_lines:
            plain_lines = [line.strip() for line in (markdown_text or "").splitlines() if line.strip()]

        for line in plain_lines:
            if line.startswith("#"):
                level = len(line) - len(line.lstrip("#"))
                text = line.lstrip("#").strip()
                heading = doc.add_heading(text, level=min(max(level, 1), 3))
                for run in heading.runs:
                    run.font.size = Pt(14)
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:].strip(), style="List Bullet")
            else:
                doc.add_paragraph(line)

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    @staticmethod
    def _html_to_lines(html: str) -> list[str]:
        text = re.sub(r"<br\s*/?>", "\n", html, flags=re.IGNORECASE)
        text = re.sub(r"</p>", "\n", text, flags=re.IGNORECASE)
        text = re.sub(r"</h([1-6])>", "\n", text, flags=re.IGNORECASE)
        text = re.sub(r"<h([1-6])[^>]*>", lambda m: "#" * int(m.group(1)) + " ", text, flags=re.IGNORECASE)
        text = re.sub(r"<li[^>]*>", "- ", text, flags=re.IGNORECASE)
        text = re.sub(r"<[^>]+>", "", text)
        lines = [line.strip() for line in text.splitlines()]
        return [line for line in lines if line]
